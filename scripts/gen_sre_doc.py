"""Generate SRE Agent Setup Guide as Word document."""
from docx import Document
from docx.shared import Pt
import os

doc = Document()

doc.add_heading('Azure SRE Agent - Setup & Configuration Guide', level=0)
doc.add_paragraph('Step-by-step guide for the Wintel Ops Automation project.')

doc.add_heading('Prerequisites', level=1)
for p in ['Azure subscription with ArcBox deployed (rg-arcbox-itpro)',
          'Owner or Contributor role on the subscription',
          'Access to https://sre.azure.com']:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('Step 1: Create the SRE Agent Instance', level=1)
doc.add_paragraph('1. Navigate to https://sre.azure.com')
doc.add_paragraph('2. Click "Create a new agent"')
doc.add_paragraph('3. Configure:')
t = doc.add_table(rows=4, cols=2); t.style = 'Light Grid Accent 1'
for r, (k, v) in enumerate([('Setting','Value'), ('Name','wintel-ops-agent'),
    ('Subscription','31adb513-7077-47bb-9567-8e9d2a462bcf'), ('Region','Sweden Central')]):
    t.rows[r].cells[0].text = k; t.rows[r].cells[1].text = v
doc.add_paragraph('4. Click "Create"')

doc.add_heading('Step 2: Grant RBAC Permissions', level=1)
doc.add_paragraph('Find the managed identity ID in Settings > Azure settings > Go to Identity. Then run:')
for cmd in [
    'az role assignment create --assignee <MANAGED_ID> --role Contributor --scope /subscriptions/31adb513-7077-47bb-9567-8e9d2a462bcf/resourceGroups/rg-arcbox-itpro',
    'az role assignment create --assignee <MANAGED_ID> --role Reader --scope /subscriptions/31adb513-7077-47bb-9567-8e9d2a462bcf/resourceGroups/rg-opsauto-sc',
    'az role assignment create --assignee <MANAGED_ID> --role "Log Analytics Reader" --scope /subscriptions/31adb513-7077-47bb-9567-8e9d2a462bcf/resourceGroups/rg-arcbox-itpro/providers/Microsoft.OperationalInsights/workspaces/law-arcbox-itpro-sc',
]:
    p = doc.add_paragraph(); r = p.add_run(cmd); r.font.name = 'Consolas'; r.font.size = Pt(8)
doc.add_paragraph('Verify: type "list my resources" in SRE Agent chat.')

doc.add_heading('Step 3: Connect Azure Monitor Alerts', level=1)
doc.add_paragraph('1. Go to Integrations > Incident platforms > Azure Monitor')
doc.add_paragraph('2. Link these alert rules:')
t = doc.add_table(rows=4, cols=2); t.style = 'Light Grid Accent 1'
for r, (k, v) in enumerate([('Alert Rule','Severity'), ('alert-heartbeat-loss','Severity 1'),
    ('alert-high-cpu','Severity 2'), ('alert-low-disk','Severity 2')]):
    t.rows[r].cells[0].text = k; t.rows[r].cells[1].text = v
doc.add_paragraph('3. Click Save')

doc.add_heading('Step 4: Create Incident Response Plans', level=1)
doc.add_paragraph('Go to Automate > Incident response > New response plan')
doc.add_heading('Plan A: Critical (Sev 0-1)', level=2)
t = doc.add_table(rows=4, cols=2); t.style = 'Light Grid Accent 1'
for r, (k, v) in enumerate([('Setting','Value'), ('Name','critical-incident-response'),
    ('Severity','Sev 0, Sev 1'), ('Run mode','Semi-autonomous')]):
    t.rows[r].cells[0].text = k; t.rows[r].cells[1].text = v
doc.add_paragraph('Instructions: "Investigate by checking server health, recent changes, correlating alerts. Use health-check and security-troubleshooting skills. Wait for approval before remediating."')

doc.add_heading('Plan B: Warning (Sev 2-3)', level=2)
t = doc.add_table(rows=4, cols=2); t.style = 'Light Grid Accent 1'
for r, (k, v) in enumerate([('Setting','Value'), ('Name','warning-incident-response'),
    ('Severity','Sev 2, Sev 3'), ('Run mode','Autonomous')]):
    t.rows[r].cells[0].text = k; t.rows[r].cells[1].text = v
doc.add_paragraph('Instructions: "Investigate and auto-remediate if safe. Create GLPI ticket with findings."')

doc.add_heading('Step 5: Upload Skills', level=1)
doc.add_paragraph('Builder > Skills > Create skill. Copy SKILL.md from repo for each:')
t = doc.add_table(rows=6, cols=3); t.style = 'Light Grid Accent 1'
skills = [('Skill Name','Description','Tools'),
    ('wintel-health-check-investigation','Health check: CPU, memory, disk, services','RunAzCliReadCommands, query-perf-trends'),
    ('security-agent-troubleshooting','Defender agent diagnosis + remediation','RunAzCliReadCommands, RunAzCliWriteCommands, query-security-alerts'),
    ('patch-validation','Pre/post patch checks, rollback','RunAzCliReadCommands, query-update-compliance'),
    ('compliance-investigation','Defender for Cloud non-compliance','RunAzCliReadCommands, query-compliance-state'),
    ('vmware-bau-operations','Snapshot cleanup, VM health','RunAzCliReadCommands, RunAzCliWriteCommands')]
for r, (a, b, c) in enumerate(skills):
    t.rows[r].cells[0].text = a; t.rows[r].cells[1].text = b; t.rows[r].cells[2].text = c

doc.add_heading('Step 6: Create Custom Tools', level=1)
doc.add_heading('Kusto Tools (Builder > Tools > Kusto)', level=2)
doc.add_paragraph('Target: law-arcbox-itpro-sc')
t = doc.add_table(rows=5, cols=2); t.style = 'Light Grid Accent 1'
for r, (k, v) in enumerate([('Tool','Source'), ('query-perf-trends','sre-tools/kusto/query-perf-trends.kql'),
    ('query-security-alerts','sre-tools/kusto/query-security-alerts.kql'),
    ('query-compliance-state','sre-tools/kusto/query-compliance-state.kql'),
    ('query-update-compliance','sre-tools/kusto/query-update-compliance.kql')]):
    t.rows[r].cells[0].text = k; t.rows[r].cells[1].text = v

doc.add_heading('Python Tools (Builder > Tools > Python)', level=2)
doc.add_paragraph('Deps: httpx, azure-cosmos')
t = doc.add_table(rows=5, cols=2); t.style = 'Light Grid Accent 1'
for r, (k, v) in enumerate([('Tool','Source'), ('glpi-create-ticket','sre-tools/python/glpi_tools.py'),
    ('glpi-query-cmdb','sre-tools/python/glpi_tools.py'),
    ('cosmos-query-runs','sre-tools/python/cosmos_tools.py'),
    ('cosmos-check-memories','sre-tools/python/cosmos_tools.py')]):
    t.rows[r].cells[0].text = k; t.rows[r].cells[1].text = v

doc.add_heading('Step 7: Build Subagents', level=1)
doc.add_paragraph('Builder > Subagent builder')
doc.add_heading('VM Diagnostics', level=2)
t = doc.add_table(rows=4, cols=2); t.style = 'Light Grid Accent 1'
for r, (k, v) in enumerate([('Setting','Value'), ('Name','vm-diagnostics'),
    ('Tools','RunAzCliReadCommands, RunAzCliWriteCommands, query-perf-trends, glpi-create-ticket, cosmos-check-memories'),
    ('Instructions','VM specialist: Check health via Arc, analyze KQL trends, check memories, root cause, create ticket if needed')]):
    t.rows[r].cells[0].text = k; t.rows[r].cells[1].text = v

doc.add_heading('Security Troubleshooter', level=2)
t = doc.add_table(rows=4, cols=2); t.style = 'Light Grid Accent 1'
for r, (k, v) in enumerate([('Setting','Value'), ('Name','security-troubleshooter'),
    ('Tools','RunAzCliReadCommands, RunAzCliWriteCommands, query-security-alerts, glpi-create-ticket'),
    ('Instructions','Security specialist: Check service, event logs, connectivity, safe remediation, escalate if fix fails')]):
    t.rows[r].cells[0].text = k; t.rows[r].cells[1].text = v
doc.add_paragraph('Test in Playground before going live.')

doc.add_heading('Step 8: Scheduled Tasks', level=1)
t = doc.add_table(rows=3, cols=3); t.style = 'Light Grid Accent 1'
for r, (a, b, c) in enumerate([('Task','Schedule','Instructions'),
    ('proactive-health-scan','Every 6h (0 */6 * * *)','Scan all Arc servers, use health-check skill for anomalies'),
    ('security-posture-check','Daily 07:00 UTC','Verify Defender agents, check new alerts in last 24h')]):
    t.rows[r].cells[0].text = a; t.rows[r].cells[1].text = b; t.rows[r].cells[2].text = c

doc.add_heading('Step 9: MCP Server (Optional)', level=1)
doc.add_paragraph('Integrations > Connectors > MCP > Add MCP server')
t = doc.add_table(rows=3, cols=2); t.style = 'Light Grid Accent 1'
for r, (k, v) in enumerate([('Setting','Value'), ('Name','glpi-itsm'),
    ('URL','http://glpi-opsauto-demo.swedencentral.azurecontainer.io')]):
    t.rows[r].cells[0].text = k; t.rows[r].cells[1].text = v

doc.add_heading('Step 10: Verify', level=1)
doc.add_heading('Test 1: Chat', level=2)
doc.add_paragraph('Ask: "What servers are in my environment?"')
doc.add_heading('Test 2: Trigger Incident', level=2)
p = doc.add_paragraph(); r = p.add_run(
    'az connectedmachine run-command create --resource-group rg-arcbox-itpro --machine-name ArcBox-Win2K22 --name stressCPU --script "while ($true) { [math]::Sqrt(12345) }" --async-execution true')
r.font.name = 'Consolas'; r.font.size = Pt(8)
doc.add_heading('Test 3: Security Issue', level=2)
p = doc.add_paragraph(); r = p.add_run(
    'az connectedmachine run-command create --resource-group rg-arcbox-itpro --machine-name ArcBox-Win2K22 --name stopDefender --script "Stop-Service -Name WinDefend -Force"')
r.font.name = 'Consolas'; r.font.size = Pt(8)

doc.add_heading('Troubleshooting', level=1)
t = doc.add_table(rows=6, cols=2); t.style = 'Light Grid Accent 1'
for r, (k, v) in enumerate([('Issue','Solution'),
    ("Can't see servers",'Check RBAC: Reader on rg-arcbox-itpro'),
    ('No incidents','Check alert action group linked to SRE Agent'),
    ('Skills not loading','Check description matches context; auto-loads'),
    ('Subagent not invoked','Use /agent vm-diagnostics explicitly'),
    ('KQL failing','Check workspace ID and Log Analytics Reader role')]):
    t.rows[r].cells[0].text = k; t.rows[r].cells[1].text = v

path = os.path.join(os.path.expanduser('~'), 'Downloads', 'SRE-Agent-Setup-Guide.docx')
doc.save(path)
print('Saved to: ' + path)
